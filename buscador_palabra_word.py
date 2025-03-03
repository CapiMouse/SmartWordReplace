from docx import Document  # Para manipular documentos Word
from docx.shared import Pt  # Para definir tamaños de fuente en puntos
# Importar el diccionario con las palabras a buscar y reemplazar
from datos_centrales import diccionario_completo, archivo_word


def reemplazar_en_runs(runs, palabras_reemplazo):
    """
    Reemplaza palabras en los `runs` de un párrafo o celda mientras conserva el formato.

    Args:
        runs (list): Lista de `runs` del párrafo o celda.
        palabras_reemplazo (dict): Diccionario con palabras y valores a reemplazar.
    """
    # Reconstruir texto completo del párrafo o celda
    texto_completo = "".join(run.text for run in runs)
    texto_modificado = texto_completo

    # Reemplazar palabras en el texto completo
    for palabra, reemplazo in palabras_reemplazo.items():
        texto_modificado = texto_modificado.replace(palabra, reemplazo)

    # Si el texto no cambió, salir
    if texto_completo == texto_modificado:
        return

    # Guardar formato original de cada run
    formatos_originales = []
    for run in runs:
        formatos_originales.append({
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "font_name": run.font.name,
            "font_size": run.font.size,
            "text": run.text
        })

    # Limpiar los textos de los `runs` actuales
    for run in runs:
        run.text = ""

    # Redistribuir el texto modificado entre los `runs`
    texto_restante = texto_modificado
    for i, run in enumerate(runs):
        if texto_restante:
            longitud_run = len(formatos_originales[i]["text"])
            run.text = texto_restante[:longitud_run]
            texto_restante = texto_restante[longitud_run:]

            # Aplicar el formato deseado al texto redistribuido
            run.font.name = "Bahnschrift"
            run.font.size = Pt(12)

            # Restaurar el formato original (negrita, cursiva, etc.)
            run.bold = formatos_originales[i]["bold"]
            run.italic = formatos_originales[i]["italic"]
            run.underline = formatos_originales[i]["underline"]

    # Si queda texto restante, añadirlo de manera segura
    if texto_restante:
        ultimo_parrafo = runs[0]._parent  # Obtener el párrafo padre
        nuevo_run = ultimo_parrafo.add_run(texto_restante)
        # Debajo elegimos el tamaño y formato del texto que vamos a sustiuir
        nuevo_run.font.name = "Bahnschrift"
        nuevo_run.font.size = Pt(12)

    print("Texto modificado y redistribuido correctamente.")


def buscar_y_reemplazar_en_word(ruta_archivo, diccionario):
    """
    Busca palabras en un documento Word (.docx) y las reemplaza por los valores del diccionario.
    Realiza la búsqueda y reemplazo en párrafos y tablas, preservando el formato original.

    Args:
        ruta_archivo (str): Ruta del archivo Word donde se realizará la búsqueda y reemplazo.

    Returns:
        None
    """
    # Cargar el documento
    try:
        doc = Document(ruta_archivo)
    except Exception as e:
        print(f"Error al abrir el archivo Word: {type(e)}")
        return

    """Aqui se encuentra el diccionario creado proviene del modulo diccionario_busqueda 
    proviene del modulo diccionario_creado.py"""

    palabras_reemplazo = diccionario

    # Reemplazar en los párrafos que contenga el word
    for parrafo in doc.paragraphs:
        reemplazar_en_runs(parrafo.runs, palabras_reemplazo)

    # Reemplazar en las tablas que contenga el word
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                # Iterar por todos los párrafos dentro de cada celda
                for parrafo in celda.paragraphs:
                    reemplazar_en_runs(parrafo.runs, palabras_reemplazo)

    # Sobrescribir el documento original
    try:
        doc.save(ruta_archivo)
        print(f"Documento sobrescrito exitosamente en: {ruta_archivo}")
    except Exception as e:
        print(f"Error al sobrescribir el archivo Word: {e}")


if __name__ == "__main__":
    buscar_y_reemplazar_en_word(archivo_word, diccionario_completo)
